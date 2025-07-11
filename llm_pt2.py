import os
import streamlit as st
import boto3
import json
import random
from datetime import datetime, timedelta
from pathlib import Path
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import base64
from io import BytesIO
from PIL import Image

# Librerías para generación de documentos
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

# ============================================================================
# CONFIGURACIÓN DE STREAMLIT
# ============================================================================

st.set_page_config(
    page_title="Sistema PQRS - Veolia Colombia",
    page_icon="💧",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Ruta del logo
LOGO_PATH = "/Users/michaelmoreno/Desktop/T Logic/Veolia/sistemapqrs/RGB_VEOLIA_HD-1024x418.webp"

# Función para cargar el logo
def get_logo_base64():
    """Convierte el logo a base64 para usar en HTML"""
    try:
        with open(LOGO_PATH, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode()
    except:
        return None

# CSS Profesional con colores de Veolia
logo_base64 = get_logo_base64()
st.markdown(f"""
<style>
    /* Variables de color Veolia */
    :root {{
        --veolia-primary: #00A982;
        --veolia-secondary: #004B87;
        --veolia-light: #E8F5F1;
        --veolia-dark: #003865;
        --veolia-accent: #7ED321;
        --text-primary: #212529;
        --text-secondary: #6C757D;
        --background: #F8F9FA;
        --card-background: #FFFFFF;
        --border-color: #DEE2E6;
        --success-color: #28A745;
        --warning-color: #FFC107;
        --error-color: #DC3545;
    }}

    /* Estilos generales */
    .stApp {{
        background-color: var(--background);
    }}

    /* Logo container */
    .logo-container {{
        display: flex;
        justify-content: center;
        align-items: center;
        padding: 1rem 0;
        background: white;
        border-radius: 10px;
        margin-bottom: 1rem;
    }}

    .logo-container img {{
        max-width: 200px;
        height: auto;
    }}

    /* Header principal con logo */
    .main-header {{
        background: linear-gradient(135deg, var(--veolia-secondary) 0%, var(--veolia-primary) 100%);
        padding: 2.5rem;
        border-radius: 20px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.15);
        position: relative;
        overflow: hidden;
    }}

    .main-header::before {{
        content: "";
        position: absolute;
        top: -50%;
        right: -50%;
        width: 200%;
        height: 200%;
        background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%);
        transform: rotate(45deg);
    }}

    .main-header-logo {{
        width: 150px;
        margin-bottom: 1rem;
        filter: brightness(0) invert(1);
    }}

    .main-header h1 {{
        font-size: 2.5rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
        position: relative;
        z-index: 1;
    }}

    .main-header p {{
        font-size: 1.1rem;
        opacity: 0.95;
        position: relative;
        z-index: 1;
    }}

    /* Cards y contenedores */
    .service-card {{
        background: var(--card-background);
        padding: 2rem;
        border-radius: 15px;
        box-shadow: 0 5px 20px rgba(0,0,0,0.08);
        margin: 1.5rem 0;
        border: 1px solid var(--border-color);
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }}

    .service-card:hover {{
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.12);
        border-color: var(--veolia-primary);
    }}

    .service-card::after {{
        content: "";
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 3px;
        background: var(--veolia-primary);
        transition: left 0.3s ease;
    }}

    .service-card:hover::after {{
        left: 0;
    }}

    /* Métricas */
    .metric-container {{
        background: linear-gradient(135deg, var(--veolia-light) 0%, #FFFFFF 100%);
        padding: 1.5rem;
        border-radius: 12px;
        text-align: center;
        box-shadow: 0 4px 15px rgba(0,0,0,0.05);
        border: 1px solid var(--veolia-primary);
        transition: all 0.3s ease;
    }}

    .metric-container:hover {{
        transform: scale(1.02);
        box-shadow: 0 6px 20px rgba(0,0,0,0.08);
    }}

    .metric-value {{
        font-size: 2rem;
        font-weight: 700;
        color: var(--veolia-secondary);
        margin: 0.5rem 0;
    }}

    .metric-label {{
        font-size: 0.9rem;
        color: var(--text-secondary);
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }}

    /* Documento preview con logo */
    .document-preview {{
        background: white;
        border: 1px solid var(--border-color);
        border-radius: 8px;
        padding: 2rem;
        margin: 1rem 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        font-family: 'Times New Roman', serif;
        line-height: 1.6;
    }}

    .document-header {{
        display: flex;
        justify-content: space-between;
        align-items: flex-start;
        margin-bottom: 2rem;
    }}

    .document-header-text {{
        text-align: right;
        color: var(--veolia-secondary);
        font-weight: bold;
    }}

    .document-logo {{
        width: 120px;
        height: auto;
    }}

    .document-subject {{
        font-weight: bold;
        margin: 1rem 0;
    }}

    /* Botones personalizados */
    .stButton > button {{
        background: var(--veolia-primary);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(0,169,130,0.3);
    }}

    .stButton > button:hover {{
        background: var(--veolia-secondary);
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(0,169,130,0.4);
    }}

    /* Selectbox y inputs */
    .stSelectbox > div > div {{
        background: white;
        border: 2px solid var(--border-color);
        border-radius: 8px;
        transition: all 0.3s ease;
    }}

    .stSelectbox > div > div:focus-within {{
        border-color: var(--veolia-primary);
        box-shadow: 0 0 0 3px rgba(0,169,130,0.1);
    }}

    .stTextInput > div > div {{
        background: white;
        border: 2px solid var(--border-color);
        border-radius: 8px;
        transition: all 0.3s ease;
    }}

    .stTextInput > div > div:focus-within {{
        border-color: var(--veolia-primary);
        box-shadow: 0 0 0 3px rgba(0,169,130,0.1);
    }}

    /* Tabs mejoradas */
    .stTabs [data-baseweb="tab-list"] {{
        gap: 2rem;
        background: transparent;
        border-bottom: 2px solid var(--border-color);
    }}

    .stTabs [data-baseweb="tab"] {{
        height: 50px;
        padding: 0 2rem;
        background: transparent;
        border: none;
        color: var(--text-secondary);
        font-weight: 600;
        transition: all 0.3s ease;
    }}

    .stTabs [aria-selected="true"] {{
        background: transparent;
        color: var(--veolia-primary);
        border-bottom: 3px solid var(--veolia-primary);
    }}

    /* Sidebar */
    .css-1d391kg {{
        background: var(--card-background);
        padding: 1.5rem;
    }}

    /* Info boxes */
    .info-box {{
        background: var(--veolia-light);
        border-left: 4px solid var(--veolia-primary);
        padding: 1rem 1.5rem;
        border-radius: 0 8px 8px 0;
        margin: 1rem 0;
    }}

    .warning-box {{
        background: #FFF3CD;
        border-left: 4px solid var(--warning-color);
        padding: 1rem 1.5rem;
        border-radius: 0 8px 8px 0;
        margin: 1rem 0;
    }}

    .success-box {{
        background: #D4EDDA;
        border-left: 4px solid var(--success-color);
        padding: 1rem 1.5rem;
        border-radius: 0 8px 8px 0;
        margin: 1rem 0;
    }}

    /* Estados de PQRS */
    .pqrs-type-badge {{
        display: inline-block;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-weight: 600;
        font-size: 0.9rem;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }}

    .pqrs-peticion {{
        background: #E3F2FD;
        color: #1976D2;
    }}

    .pqrs-queja {{
        background: #FFF3E0;
        color: #F57C00;
    }}

    .pqrs-reclamo {{
        background: #FFEBEE;
        color: #D32F2F;
    }}

    .pqrs-sugerencia {{
        background: #E8F5E9;
        color: #388E3C;
    }}

    /* Animaciones */
    @keyframes fadeIn {{
        from {{ opacity: 0; transform: translateY(10px); }}
        to {{ opacity: 1; transform: translateY(0); }}
    }}

    .fade-in {{
        animation: fadeIn 0.5s ease-out;
    }}

    /* Progress indicator */
    .progress-step {{
        display: flex;
        align-items: center;
        margin: 1rem 0;
    }}

    .progress-circle {{
        width: 40px;
        height: 40px;
        border-radius: 50%;
        background: var(--veolia-primary);
        color: white;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: bold;
        margin-right: 1rem;
    }}

    .progress-circle.inactive {{
        background: var(--border-color);
        color: var(--text-secondary);
    }}

    /* Responsive */
    @media (max-width: 768px) {{
        .main-header h1 {{
            font-size: 2rem;
        }}
        
        .service-card {{
            padding: 1.5rem;
        }}
        
        .document-logo {{
            width: 80px;
        }}
    }}
</style>
""", unsafe_allow_html=True)

# ============================================================================
# CLASES PRINCIPALES
# ============================================================================

class VeoliaPQRSGenerator:
    """Generador de respuestas PQRS usando AWS Bedrock"""
    
    def __init__(self):
        self.bedrock_client = None
        self.model_id = "us.anthropic.claude-sonnet-4-20250514-v1:0"
        self.tipos_pqrs = {
            'P': 'PETICIÓN',
            'Q': 'QUEJA',
            'R': 'RECLAMO',
            'S': 'SUGERENCIA'
        }
        self.logo_path = LOGO_PATH
        
        # Inicializar Bedrock
        try:
            self.bedrock_client = boto3.client(
                service_name='bedrock-runtime',
                region_name='us-east-1',
                aws_access_key_id=st.secrets.get("AWS_ACCESS_KEY_ID"),
                aws_secret_access_key=st.secrets.get("AWS_SECRET_ACCESS_KEY")
            )
        except Exception as e:
            st.error(f"Error inicializando Bedrock: {e}")
    
    def generar_radicado(self, tipo_pqrs):
        """Genera un número de radicado único"""
        fecha = datetime.now().strftime("%Y%m%d")
        numero = datetime.now().strftime("%H%M%S")
        return f"VEO-{tipo_pqrs}-{fecha}-{numero}"
    
    def generar_datos_cliente(self, numero_contrato):
        """Genera datos ficticios del cliente basados en el número de contrato"""
        random.seed(int(numero_contrato))
        
        barrios = ["Santa Bárbara", "Chapinero Alto", "Usaquén", "Cedritos", "La Castellana", 
                   "Salitre", "Teusaquillo", "Chicó", "Rosales", "Colina Campestre"]
        
        nombres = ["Ana María", "Juan Carlos", "Patricia", "Luis Alberto", "Carolina", 
                  "José Manuel", "Martha Lucía", "Carlos Andrés", "María José", "Diego Alejandro"]
        
        apellidos = ["González", "Rodríguez", "Martínez", "López", "Sánchez", 
                    "Ramírez", "Torres", "Herrera", "Jiménez", "Morales"]
        
        # Generar datos del cliente
        nombre = random.choice(nombres)
        apellido1 = random.choice(apellidos)
        apellido2 = random.choice(apellidos)
        nombre_completo = f"{nombre} {apellido1} {apellido2}"
        
        # Generar consumos históricos
        consumo_base = random.randint(15, 35)
        consumos_historicos = []
        meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio"]
        
        for i in range(6):
            variacion = random.randint(-3, 3)
            consumo = max(10, consumo_base + variacion)
            consumos_historicos.append({
                'mes': meses[i],
                'consumo': consumo,
                'año': 2024
            })
        
        promedio_consumo = sum(c['consumo'] for c in consumos_historicos) / len(consumos_historicos)
        ultimo_consumo = consumos_historicos[-1]['consumo']
        
        # Calcular valores según estrato
        estrato = random.randint(1, 6)
        tarifas_m3 = {1: 2000, 2: 2500, 3: 3000, 4: 3500, 5: 4000, 6: 4500}
        valor_m3 = tarifas_m3[estrato]
        
        # Generar correo
        nombre_limpio = nombre.lower().replace(' ', '').replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
        apellido_limpio = apellido1.lower().replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
        correo = f"{nombre_limpio}.{apellido_limpio}@gmail.com"
        
        return {
            "numero_contrato": numero_contrato,
            "nombre_completo": nombre_completo,
            "cedula": f"{random.randint(80000000, 99999999)}",
            "direccion": f"Calle {random.randint(1, 150)} #{random.randint(1, 99)}-{random.randint(1, 99)}, {random.choice(barrios)}",
            "correo": correo,
            "telefono": f"3{random.randint(10, 50)}{random.randint(1000000, 9999999)}",
            "estrato": estrato,
            "consumo_actual": ultimo_consumo,
            "consumo_promedio": round(promedio_consumo, 1),
            "consumos_historicos": consumos_historicos,
            "valor_m3": valor_m3,
            "valor_factura": round(ultimo_consumo * valor_m3, -2),
            "fecha_ultima_lectura": (datetime.now() - timedelta(days=random.randint(1, 15))).strftime('%Y-%m-%d'),
            "numero_medidor": f"MED-{random.randint(10000, 99999)}",
            "tipo_usuario": random.choice(["Residencial", "Comercial"]),
            "fecha_instalacion": (datetime.now() - timedelta(days=random.randint(365, 3650))).strftime('%Y-%m-%d'),
            "barrio": random.choice(barrios),
            "ciclo_facturacion": random.randint(1, 6)
        }
    
    def generar_contexto_pqrs(self, tipo_pqrs, datos_cliente):
        """Genera el contexto específico según el tipo de PQRS"""
        contextos = {
            'P': f"""El usuario solicita información detallada sobre su cuenta de servicios con número de contrato {datos_cliente['numero_contrato']}. 
                    Requiere conocer el histórico de consumos de los últimos 6 meses, las tarifas aplicadas según su estrato {datos_cliente['estrato']},
                    y aclaración sobre los componentes de la factura. También solicita información sobre programas de ahorro de agua disponibles.""",
            
            'Q': f"""El usuario presenta una queja formal por la atención recibida durante la visita técnica realizada el {datos_cliente['fecha_ultima_lectura']} 
                    en su predio ubicado en {datos_cliente['direccion']}. El técnico no siguió los protocolos de servicio, no presentó identificación 
                    y dejó el área de trabajo en desorden. Solicita medidas correctivas y una nueva visita técnica.""",
            
            'R': f"""El usuario presenta un reclamo por el alto consumo facturado en el último periodo. El consumo registrado de {datos_cliente['consumo_actual']} m³ 
                    es significativamente mayor al promedio histórico de {datos_cliente['consumo_promedio']} m³. El valor facturado de ${datos_cliente['valor_factura']:,} 
                    no corresponde con el patrón de consumo habitual. Solicita revisión técnica del medidor {datos_cliente['numero_medidor']} y ajuste en la factura.""",
            
            'S': f"""El usuario, cliente desde {datos_cliente['fecha_instalacion']}, sugiere implementar mejoras en el sistema de notificación de lecturas 
                    y en la aplicación móvil. Propone incluir alertas de consumo inusual, gráficas comparativas mensuales y la opción de programar 
                    visitas técnicas directamente desde la app. También sugiere implementar un sistema de puntos por ahorro de agua."""
        }
        
        return contextos.get(tipo_pqrs, "")
    
    def generar_respuesta_bedrock(self, tipo_pqrs, datos_cliente):
        """Genera la respuesta usando Claude a través de Bedrock"""
        if not self.bedrock_client:
            return None, None, "Bedrock no está configurado correctamente"
        
        radicado = self.generar_radicado(tipo_pqrs)
        contexto = self.generar_contexto_pqrs(tipo_pqrs, datos_cliente)
        
        # Construir el prompt
        system_prompt = """Eres un representante experto del servicio al cliente de Veolia Colombia, 
        empresa líder en gestión del agua y servicios ambientales. Tu rol es generar respuestas 
        profesionales, empáticas y completas a las PQRS de los usuarios.
        
        Debes mantener un tono profesional pero cercano, demostrar conocimiento técnico cuando 
        sea necesario y siempre expresar el compromiso de Veolia con la calidad del servicio 
        y el cuidado del medio ambiente."""
        
        user_prompt = f"""Genera una respuesta formal y completa para la siguiente {self.tipos_pqrs[tipo_pqrs]}:

Radicado: {radicado}
Fecha: {datetime.now().strftime('%d de %B de %Y')}

Datos del usuario:
- Nombre: {datos_cliente['nombre_completo']}
- Cédula: {datos_cliente['cedula']}
- Contrato: {datos_cliente['numero_contrato']}
- Dirección: {datos_cliente['direccion']}
- Estrato: {datos_cliente['estrato']}
- Tipo de usuario: {datos_cliente['tipo_usuario']}

Contexto de la {self.tipos_pqrs[tipo_pqrs]}:
{contexto}

La respuesta debe:
1. Iniciar con un saludo cordial y acuse de recibo
2. Abordar específicamente todos los puntos planteados
3. Proporcionar información técnica cuando sea relevante
4. Mencionar el marco legal aplicable (Ley 142 de 1994, Resoluciones CRA)
5. Detallar los pasos a seguir y tiempos de respuesta
6. Incluir información de contacto y canales de atención
7. Cerrar con un mensaje de compromiso con el servicio
8. Mantener un formato de carta formal pero con lenguaje claro y cercano

Genera la respuesta completa sin usar títulos ni numeraciones, manteniendo un flujo natural."""
        
        try:
            # Configurar la solicitud
            body = json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 3000,
                "system": system_prompt,
                "messages": [
                    {
                        "role": "user",
                        "content": user_prompt
                    }
                ],
                "temperature": 0.7,
                "top_p": 0.9,
            })
            
            # Invocar el modelo
            response = self.bedrock_client.invoke_model(
                body=body,
                modelId=self.model_id,
                accept='application/json',
                contentType='application/json'
            )
            
            # Procesar respuesta
            response_body = json.loads(response.get('body').read())
            respuesta_texto = response_body['content'][0]['text'].strip()
            
            return respuesta_texto, radicado, None
            
        except Exception as e:
            return None, None, f"Error generando respuesta: {str(e)}"
    
    def generar_documento_word(self, texto, radicado, datos_cliente):
        """Genera documento Word con la respuesta"""
        doc = Document()
        
        # Configurar márgenes
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)
        
        # Agregar logo en el encabezado
        header = doc.sections[0].header
        header_table = header.add_table(rows=1, cols=2, width=doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)
        header_table.autofit = False
        
        # Celda del logo
        logo_cell = header_table.cell(0, 0)
        logo_paragraph = logo_cell.paragraphs[0]
        
        try:
            run = logo_paragraph.add_run()
            run.add_picture(self.logo_path, width=Inches(2))
        except:
            # Si no se puede cargar el logo, agregar texto
            run = logo_paragraph.add_run('VEOLIA')
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 75, 135)
        
        # Celda del texto
        text_cell = header_table.cell(0, 1)
        text_paragraph = text_cell.paragraphs[0]
        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = text_paragraph.add_run('Gestión del Agua y Servicios Ambientales')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 75, 135)
        
        # Espaciado después del encabezado
        doc.add_paragraph()
        
        # Fecha y radicado
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Bogotá D.C., {datetime.now().strftime('%d de %B de %Y')}\n").font.size = Pt(11)
        p.add_run(f"Radicado: {radicado}").font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Destinatario
        p = doc.add_paragraph()
        p.add_run(f"Señor(a)\n{datos_cliente['nombre_completo']}\n{datos_cliente['direccion']}\nBogotá D.C.").font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Asunto
        p = doc.add_paragraph()
        run = p.add_run(f"Asunto: Respuesta a {self.tipos_pqrs[radicado[4]].lower()} radicada")
        run.font.bold = True
        run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Saludo
        p = doc.add_paragraph()
        p.add_run("Respetado(a) señor(a):").font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Contenido
        for parrafo in texto.split('\n\n'):
            if parrafo.strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(parrafo).font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Despedida
        p = doc.add_paragraph()
        p.add_run("Cordialmente,").font.size = Pt(11)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Firma
        p = doc.add_paragraph()
        run = p.add_run("MARÍA FERNANDA LÓPEZ GARCÍA\nCoordinadora Servicio al Cliente\nVeolia Colombia")
        run.font.bold = True
        run.font.size = Pt(11)
        
        # Pie de página
        footer = doc.sections[0].footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Veolia Colombia - Comprometidos con el Medio Ambiente\n')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0, 169, 130)
        run = p.add_run('Línea gratuita nacional: 01 8000 123 456 - www.veolia.com.co')
        run.font.size = Pt(8)
        
        # Guardar en memoria
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer

# ============================================================================
# FUNCIONES DE UTILIDAD
# ============================================================================

def mostrar_datos_cliente(datos):
    """Muestra los datos del cliente en un formato visual atractivo"""
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="metric-container">
            <div class="metric-label">Contrato</div>
            <div class="metric-value">{}</div>
        </div>
        """.format(datos['numero_contrato']), unsafe_allow_html=True)
        
        st.markdown("""
        <div class="info-box">
            <strong>Cliente:</strong><br>
            {}<br>
            CC: {}
        </div>
        """.format(datos['nombre_completo'], datos['cedula']), unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="metric-container">
            <div class="metric-label">Consumo Actual</div>
            <div class="metric-value">{} m³</div>
        </div>
        """.format(datos['consumo_actual']), unsafe_allow_html=True)
        
        st.markdown("""
        <div class="info-box">
            <strong>Dirección:</strong><br>
            {}<br>
            Estrato: {}
        </div>
        """.format(datos['direccion'], datos['estrato']), unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="metric-container">
            <div class="metric-label">Valor Factura</div>
            <div class="metric-value">${:,}</div>
        </div>
        """.format(datos['valor_factura']), unsafe_allow_html=True)
        
        st.markdown("""
        <div class="info-box">
            <strong>Medidor:</strong> {}<br>
            <strong>Tipo:</strong> {}
        </div>
        """.format(datos['numero_medidor'], datos['tipo_usuario']), unsafe_allow_html=True)

def crear_grafica_consumos(datos):
    """Crea gráfica de consumos históricos"""
    df_consumos = pd.DataFrame(datos['consumos_historicos'])
    
    fig = go.Figure()
    
    # Línea de consumo
    fig.add_trace(go.Scatter(
        x=df_consumos['mes'],
        y=df_consumos['consumo'],
        mode='lines+markers',
        name='Consumo mensual',
        line=dict(color='#00A982', width=3),
        marker=dict(size=10, color='#00A982')
    ))
    
    # Línea de promedio
    promedio = datos['consumo_promedio']
    fig.add_trace(go.Scatter(
        x=df_consumos['mes'],
        y=[promedio] * len(df_consumos),
        mode='lines',
        name='Promedio',
        line=dict(color='#004B87', width=2, dash='dash')
    ))
    
    fig.update_layout(
        title='Histórico de Consumos (m³)',
        xaxis_title='Mes',
        yaxis_title='Consumo (m³)',
        height=400,
        hovermode='x unified',
        showlegend=True,
        plot_bgcolor='white',
        paper_bgcolor='white'
    )
    
    fig.update_xaxes(gridcolor='lightgray')
    fig.update_yaxes(gridcolor='lightgray')
    
    return fig

def get_tipo_badge(tipo):
    """Retorna el HTML para el badge del tipo de PQRS"""
    badges = {
        'P': '<span class="pqrs-type-badge pqrs-peticion">PETICIÓN</span>',
        'Q': '<span class="pqrs-type-badge pqrs-queja">QUEJA</span>',
        'R': '<span class="pqrs-type-badge pqrs-reclamo">RECLAMO</span>',
        'S': '<span class="pqrs-type-badge pqrs-sugerencia">SUGERENCIA</span>'
    }
    return badges.get(tipo, '')

# ============================================================================
# INTERFAZ PRINCIPAL
# ============================================================================

# Header principal con logo
if logo_base64:
    st.markdown(f"""
    <div class="main-header">
        <img src="data:image/webp;base64,{logo_base64}" class="main-header-logo" alt="Veolia Logo">
        <h1>💧 Sistema de Gestión PQRS</h1>
        <p>Atención al Cliente</p>
    </div>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <div class="main-header">
        <h1>💧 Sistema de Gestión PQRS</h1>
        <p>Veolia Colombia - Atención al Cliente</p>
    </div>
    """, unsafe_allow_html=True)

# Inicializar generador
generador = VeoliaPQRSGenerator()

# Sidebar con logo
with st.sidebar:
    # Mostrar logo en el sidebar si existe
    if os.path.exists(LOGO_PATH):
        st.image(LOGO_PATH, width=200)
        st.markdown("---")
    
    st.markdown("### 🔧 Panel de Control")
    
    # Estado del sistema
    bedrock_status = "✅ Activo" if generador.bedrock_client else "❌ Inactivo"
    
    st.markdown(f"""
    <div class="service-card">
        <h4>Estado del Sistema</h4>
        <p><strong>AWS Bedrock:</strong> {bedrock_status}</p>
        <p><strong>Modelo:</strong> Claude Sonnet</p>
        <p><strong>Región:</strong> us-east-1</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Estadísticas del día
    st.markdown("""
    <div class="service-card">
        <h4>📊 Estadísticas del Día</h4>
        <p><strong>PQRS Generadas:</strong> 12</p>
        <p><strong>Tiempo Promedio:</strong> 2.3 min</p>
        <p><strong>Satisfacción:</strong> 98%</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Información adicional
    with st.expander("ℹ️ Acerca del Sistema"):
        st.markdown("""
        **Sistema PQRS Veolia v2.0**
        
        Este sistema utiliza inteligencia artificial para:
        - Generar respuestas personalizadas
        - Cumplir con normativa vigente
        - Mantener estándares de calidad
        - Reducir tiempos de respuesta
        
        **Normativa aplicable:**
        - Ley 142 de 1994
        - Resolución CRA 413 de 2006
        - Decreto 1077 de 2015
        """)

# Contenido principal
tabs = st.tabs(["📝 Generar PQRS", "📊 Dashboard", "📚 Historial", "❓ Ayuda"])

# Tab 1: Generar PQRS
with tabs[0]:
    st.markdown("### Generador de Respuestas PQRS")
    
    # Progreso visual
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("""
        <div class="progress-step">
            <div class="progress-circle">1</div>
            <div>Seleccionar Tipo</div>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown("""
        <div class="progress-step">
            <div class="progress-circle">2</div>
            <div>Ingresar Contrato</div>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown("""
        <div class="progress-step">
            <div class="progress-circle inactive">3</div>
            <div>Generar Respuesta</div>
        </div>
        """, unsafe_allow_html=True)
    
    # Formulario principal
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown('<div class="service-card">', unsafe_allow_html=True)
        st.markdown("#### 1️⃣ Tipo de PQRS")
        
        tipo_pqrs = st.selectbox(
            "Seleccione el tipo",
            options=['P', 'Q', 'R', 'S'],
            format_func=lambda x: generador.tipos_pqrs[x],
            help="Seleccione el tipo de solicitud a procesar"
        )
        
        st.markdown(get_tipo_badge(tipo_pqrs), unsafe_allow_html=True)
        
        # Descripción del tipo seleccionado
        descripciones = {
            'P': "📋 **Petición:** Solicitud de información o documentos",
            'Q': "😠 **Queja:** Inconformidad con el servicio o atención",
            'R': "💰 **Reclamo:** Inconformidad con facturación o cobros",
            'S': "💡 **Sugerencia:** Propuestas de mejora del servicio"
        }
        
        st.markdown(f"\n{descripciones[tipo_pqrs]}")
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="service-card">', unsafe_allow_html=True)
        st.markdown("#### 2️⃣ Datos del Cliente")
        
        numero_contrato = st.text_input(
            "Número de Contrato",
            placeholder="Ej: 1234567890",
            help="Ingrese el número de contrato del cliente (10 dígitos)"
        )
        
        if numero_contrato and len(numero_contrato) == 10 and numero_contrato.isdigit():
            # Generar y mostrar datos del cliente
            datos_cliente = generador.generar_datos_cliente(numero_contrato)
            
            st.success(f"✅ Cliente encontrado: {datos_cliente['nombre_completo']}")
            
            # Botón para generar respuesta
            if st.button("🚀 Generar Respuesta PQRS", type="primary", use_container_width=True):
                with st.spinner("Generando respuesta con IA..."):
                    respuesta, radicado, error = generador.generar_respuesta_bedrock(tipo_pqrs, datos_cliente)
                    
                    if error:
                        st.error(f"❌ {error}")
                    else:
                        st.session_state.ultima_respuesta = {
                            'texto': respuesta,
                            'radicado': radicado,
                            'datos_cliente': datos_cliente,
                            'tipo': tipo_pqrs
                        }
                        st.success(f"✅ Respuesta generada - Radicado: {radicado}")
        elif numero_contrato:
            st.warning("⚠️ El número de contrato debe tener exactamente 10 dígitos")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Mostrar respuesta generada
    if 'ultima_respuesta' in st.session_state:
        st.markdown("---")
        st.markdown("### 📄 Respuesta Generada")
        
        respuesta_data = st.session_state.ultima_respuesta
        
        # Información del radicado
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(f"**Radicado:** {respuesta_data['radicado']}")
        with col2:
            st.markdown(f"**Tipo:** {get_tipo_badge(respuesta_data['tipo'])}", unsafe_allow_html=True)
        with col3:
            st.markdown(f"**Fecha:** {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        # Datos del cliente
        st.markdown("#### Información del Cliente")
        mostrar_datos_cliente(respuesta_data['datos_cliente'])
        
        # Gráfica de consumos
        st.markdown("#### Histórico de Consumos")
        fig = crear_grafica_consumos(respuesta_data['datos_cliente'])
        st.plotly_chart(fig, use_container_width=True)
        
        # Vista previa del documento con logo
        st.markdown("#### Vista Previa de la Respuesta")
        
        if logo_base64:
            st.markdown(f"""
            <div class="document-preview">
                <div class="document-header">
                    <img src="data:image/webp;base64,{logo_base64}" class="document-logo" alt="Veolia">
                    <div class="document-header-text">
                        Gestión del Agua y Servicios Ambientales
                    </div>
                </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown('<div class="document-preview">', unsafe_allow_html=True)
            st.markdown(f"""
            <div class="document-header">
                <div></div>
                <div class="document-header-text">
                    VEOLIA COLOMBIA<br>
                    Gestión del Agua y Servicios Ambientales
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown(f"""
        <p style="text-align: right;">
            Bogotá D.C., {datetime.now().strftime('%d de %B de %Y')}<br>
            Radicado: {respuesta_data['radicado']}
        </p>
        
        <p>
            Señor(a)<br>
            {respuesta_data['datos_cliente']['nombre_completo']}<br>
            {respuesta_data['datos_cliente']['direccion']}<br>
            Bogotá D.C.
        </p>
        
        <p class="document-subject">
            Asunto: Respuesta a {generador.tipos_pqrs[respuesta_data['tipo']].lower()} radicada
        </p>
        
        <p>Respetado(a) señor(a):</p>
        """, unsafe_allow_html=True)
        
        # Mostrar contenido de la respuesta
        for parrafo in respuesta_data['texto'].split('\n\n'):
            if parrafo.strip():
                st.markdown(f"<p style='text-align: justify;'>{parrafo}</p>", unsafe_allow_html=True)
        
        st.markdown("""
        <p>Cordialmente,</p>
        <br><br>
        <p>
            <strong>MARÍA FERNANDA LÓPEZ GARCÍA</strong><br>
            Coordinadora Servicio al Cliente<br>
            Veolia Colombia
        </p>
        """, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Botones de descarga
        st.markdown("#### Opciones de Descarga")
        col1, col2 = st.columns(2)
        
        with col1:
            # Generar documento Word
            doc_buffer = generador.generar_documento_word(
                respuesta_data['texto'],
                respuesta_data['radicado'],
                respuesta_data['datos_cliente']
            )
            
            st.download_button(
                label="📄 Descargar Word",
                data=doc_buffer.getvalue(),
                file_name=f"{respuesta_data['radicado']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with col2:
            st.button("📧 Enviar por Email", type="secondary", use_container_width=True)

# Tab 2: Dashboard
with tabs[1]:
    st.markdown("### 📊 Dashboard de PQRS")
    
    # Métricas principales
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("""
        <div class="metric-container">
            <div class="metric-label">PQRS Hoy</div>
            <div class="metric-value">28</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="metric-container">
            <div class="metric-label">Tiempo Promedio</div>
            <div class="metric-value">2.3 min</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="metric-container">
            <div class="metric-label">Resueltas</div>
            <div class="metric-value">96%</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown("""
        <div class="metric-container">
            <div class="metric-label">Satisfacción</div>
            <div class="metric-value">4.8/5</div>
        </div>
        """, unsafe_allow_html=True)
    
    # Gráficas
    col1, col2 = st.columns(2)
    
    with col1:
        # Distribución por tipo
        tipos_data = pd.DataFrame({
            'Tipo': ['Peticiones', 'Quejas', 'Reclamos', 'Sugerencias'],
            'Cantidad': [45, 28, 35, 12],
            'Color': ['#1976D2', '#F57C00', '#D32F2F', '#388E3C']
        })
        
        fig = px.pie(tipos_data, values='Cantidad', names='Tipo', 
                     title='Distribución por Tipo de PQRS',
                     color_discrete_map=dict(zip(tipos_data['Tipo'], tipos_data['Color'])))
        fig.update_traces(textposition='inside', textinfo='percent+label')
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        # Tendencia mensual
        meses = ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun']
        valores = [120, 135, 128, 142, 155, 148]
        
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=meses, y=valores,
            mode='lines+markers',
            name='PQRS Mensuales',
            line=dict(color='#00A982', width=3),
            marker=dict(size=10)
        ))
        
        fig.update_layout(
            title='Tendencia Mensual de PQRS',
            xaxis_title='Mes',
            yaxis_title='Cantidad',
            showlegend=False,
            plot_bgcolor='white',
            paper_bgcolor='white'
        )
        
        fig.update_xaxes(gridcolor='lightgray')
        fig.update_yaxes(gridcolor='lightgray')
        
        st.plotly_chart(fig, use_container_width=True)
    
    # Tabla de últimas PQRS
    st.markdown("### 📋 Últimas PQRS Procesadas")
    
    ultimas_pqrs = pd.DataFrame({
        'Radicado': ['VEO-R-20241209-143522', 'VEO-P-20241209-142815', 'VEO-Q-20241209-141203'],
        'Tipo': ['Reclamo', 'Petición', 'Queja'],
        'Cliente': ['Ana María González', 'Juan Carlos López', 'Patricia Ramírez'],
        'Estado': ['Completada', 'Completada', 'En proceso'],
        'Tiempo': ['2.1 min', '1.8 min', '3.2 min']
    })
    
    st.dataframe(ultimas_pqrs, use_container_width=True, hide_index=True)

# Tab 3: Historial
with tabs[2]:
    st.markdown("### 📚 Historial de PQRS")
    
    # Filtros
    col1, col2, col3 = st.columns(3)
    
    with col1:
        fecha_inicio = st.date_input("Fecha Inicio", datetime.now() - timedelta(days=30))
    
    with col2:
        fecha_fin = st.date_input("Fecha Fin", datetime.now())
    
    with col3:
        tipo_filtro = st.multiselect(
            "Tipo de PQRS",
            options=['Petición', 'Queja', 'Reclamo', 'Sugerencia'],
            default=['Petición', 'Queja', 'Reclamo', 'Sugerencia']
        )
    
    # Simulación de datos históricos
    st.markdown("#### Resultados de la búsqueda")
    
    # Generar datos de ejemplo de forma más simple
    tipos_ciclo = ['R', 'P', 'Q', 'S', 'R', 'P'] * 5  # 30 elementos
    radicados = []
    for i in range(30):
        tipo = tipos_ciclo[i]
        dia = (i % 30) + 1
        hora = 8 + (i % 10)
        minuto = (i * 2) % 60
        segundo = (i * 3) % 60
        radicado = f'VEO-{tipo}-202412{dia:02d}-{hora:02d}{minuto:02d}{segundo:02d}'
        radicados.append(radicado)
    
    historico_data = pd.DataFrame({
        'Radicado': radicados,
        'Tipo': ['Reclamo', 'Petición', 'Queja', 'Sugerencia', 'Reclamo', 'Petición'] * 5,
        'Fecha': pd.date_range(start='2024-12-01', periods=30, freq='D'),
        'Cliente': [f'Cliente {i}' for i in range(1, 31)],
        'Contrato': [f'{1234567890 + i}' for i in range(30)],
        'Estado': ['Completada'] * 28 + ['En proceso'] * 2,
        'Tiempo Respuesta': [f'{random.uniform(1.5, 4.0):.1f} min' for _ in range(30)]
    })
    
    st.dataframe(
        historico_data[historico_data['Tipo'].isin(tipo_filtro)].head(20),
        use_container_width=True,
        hide_index=True
    )
    
    # Botón de exportar
    st.button("📥 Exportar a Excel", type="secondary")

# Tab 4: Ayuda
with tabs[3]:
    st.markdown("### ❓ Centro de Ayuda")
    
    with st.expander("🔍 ¿Cómo generar una respuesta PQRS?"):
        st.markdown("""
        1. **Seleccione el tipo de PQRS:** Escoja entre Petición, Queja, Reclamo o Sugerencia
        2. **Ingrese el número de contrato:** Digite los 10 dígitos del contrato del cliente
        3. **Genere la respuesta:** Haga clic en el botón para que la IA genere la respuesta
        4. **Revise y descargue:** Verifique la respuesta y descárguela en Word
        """)
    
    with st.expander("📊 ¿Cómo interpretar las métricas?"):
        st.markdown("""
        - **VTR (View-Through Rate):** Porcentaje de visualización completa
        - **Tiempo promedio:** Tiempo desde la recepción hasta la respuesta
        - **Satisfacción:** Calificación promedio de los usuarios (1-5)
        - **Resueltas:** Porcentaje de PQRS completadas exitosamente
        """)
    
    with st.expander("⚖️ Marco legal y normativo"):
        st.markdown("""
        **Principales normas aplicables:**
        
        - **Ley 142 de 1994:** Régimen de servicios públicos domiciliarios
        - **Resolución CRA 413 de 2006:** Criterios para peticiones, quejas y recursos
        - **Decreto 1077 de 2015:** Decreto único reglamentario del sector vivienda
        - **Código de Procedimiento Administrativo:** Términos para responder
        
        **Tiempos de respuesta:**
        - Peticiones de información: 10 días hábiles
        - Quejas y reclamos: 15 días hábiles
        - Consultas: 30 días hábiles
        """)
    
    with st.expander("🆘 Soporte técnico"):
        st.markdown("""
        **¿Necesita ayuda adicional?**
        
        📧 **Email:** soporte.sistemas@veolia.com.co  
        📞 **Teléfono:** (601) 756 3000 Ext. 2345  
        💬 **Chat interno:** Disponible en horario laboral  
        
        **Horario de atención:**  
        Lunes a Viernes: 7:00 AM - 6:00 PM  
        Sábados: 8:00 AM - 12:00 PM
        """)

# Footer con logo
if logo_base64:
    st.markdown(f"""
    <div style="text-align: center; padding: 2rem; color: #6C757D; margin-top: 3rem;">
        <img src="data:image/webp;base64,{logo_base64}" style="width: 100px; margin-bottom: 1rem;">
        <p>Sistema PQRS Veolia Colombia v2.0 - Powered by AWS Bedrock & Claude AI</p>
        <p>© 2024 Veolia Colombia - Todos los derechos reservados</p>
    </div>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <div style="text-align: center; padding: 2rem; color: #6C757D; margin-top: 3rem;">
        <p>Sistema PQRS Veolia Colombia v2.0 - Powered by AWS Bedrock & Claude AI</p>
        <p>© 2025 Veolia Colombia - Todos los derechos reservados</p>
    </div>
    """, unsafe_allow_html=True)